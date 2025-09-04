from flask import Blueprint, render_template, redirect, url_for, flash, current_app, send_file
import os
import json
import tempfile
import shutil
from flask_login import current_user, login_required
import datetime as dt

status_bp = Blueprint('status', __name__)

@status_bp.route('/<submission_id>')
@login_required
def view_status(submission_id):
    """View a specific submission with auto-download"""
    from models import Report, SATReport

    # Check if submission_id is valid
    if not submission_id or submission_id == 'None':
        flash('Invalid submission ID.', 'error')
        return redirect(url_for('dashboard.home'))

    report = Report.query.filter_by(id=submission_id).first()
    if not report:
        flash('Report not found.', 'error')
        return redirect(url_for('dashboard.home'))

    sat_report = SATReport.query.filter_by(report_id=report.id).first()
    if not sat_report:
        flash('Report data not found.', 'error')
        return redirect(url_for('dashboard.home'))

    try:
        stored_data = json.loads(sat_report.data_json) if sat_report.data_json else {}
    except json.JSONDecodeError:
        stored_data = {}

    approvals = json.loads(report.approvals_json) if report.approvals_json else []

    # Determine overall status
    statuses = [a.get("status", "pending") for a in approvals]
    if "rejected" in statuses:
        overall_status = "rejected"
    elif all(status == "approved" for status in statuses):
        overall_status = "approved"
    elif any(status == "approved" for status in statuses):
        overall_status = "partially_approved"
    else:
        overall_status = "pending"

    # Get submission data context with fallbacks
    submission_data = stored_data.get("context", {})
    if not submission_data:
        submission_data = stored_data  # Fallback if context doesn't exist

    # Check if report files exist
    pdf_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.pdf')
    docx_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')

    download_available = os.path.exists(pdf_path) or os.path.exists(docx_path)
    has_pdf = os.path.exists(pdf_path)

    # Determine if current user can edit this report
    can_edit = False
    if current_user.role == 'Admin':
        can_edit = True  # Admin can edit any report
    elif current_user.role == 'Engineer' and current_user.email == report.user_email:
        # Engineers can edit their own reports until approved by Automation Manager
        tm_approved = any(a.get("status") == "approved" and a.get("stage") == 1 for a in approvals)
        can_edit = not tm_approved
    elif current_user.role == 'Automation Manager':
        # Automation Manager can edit reports until approved by PM
        pm_approved = any(a.get("status") == "approved" and a.get("stage") == 2 for a in approvals)
        can_edit = not pm_approved

    # Build context similar to old version
    context = {
        "submission_id": submission_id,
        "submission_data": submission_data,
        "approvals": approvals,
        "locked": report.locked,
        "can_edit": can_edit,
        "created_at": report.created_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.created_at, dt.datetime) else report.created_at,
        "updated_at": report.updated_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.updated_at, dt.datetime) else report.updated_at,
        "user_email": report.user_email,
        "document_title": submission_data.get("DOCUMENT_TITLE", "SAT Report"),
        "project_reference": submission_data.get("PROJECT_REFERENCE", ""),
        "client_name": submission_data.get("CLIENT_NAME", ""),
        "prepared_by": submission_data.get("PREPARED_BY", ""),
        "overall_status": overall_status,
        "download_available": download_available,
        "has_pdf": has_pdf,
        "auto_download": True
    }

    return render_template('status.html', **context)

@status_bp.route('/download/<submission_id>')
@login_required
def download_report(submission_id):
    """Download the generated report"""
    try:
        # Validate submission ID
        if not submission_id or submission_id == 'None':
            current_app.logger.error(f"Invalid submission ID: {submission_id}")
            flash('Invalid submission ID.', 'error')
            return redirect(url_for('dashboard.home'))

        # FORCE REGENERATION - Skip existing file check to create fresh clean document
        permanent_path = os.path.join(current_app.config['OUTPUT_DIR'], f'SAT_Report_{submission_id}_Final.docx')
        
        # Remove existing file if it exists to force fresh generation
        if os.path.exists(permanent_path):
            try:
                os.remove(permanent_path)
                current_app.logger.info(f"Removed existing file to force fresh generation: {permanent_path}")
            except Exception as e:
                current_app.logger.warning(f"Could not remove existing file: {e}")
        
        if False:  # DISABLED - Always regenerate for now
            current_app.logger.info(f"Found existing report file: {permanent_path}")
            try:
                # Try to get document title from database, but don't fail if database is down
                from models import Report, SATReport
                report = Report.query.filter_by(id=submission_id).first()
                if report:
                    sat_report = SATReport.query.filter_by(report_id=submission_id).first()
                    if sat_report and sat_report.data_json:
                        stored_data = json.loads(sat_report.data_json)
                        context_data = stored_data.get("context", {})
                        doc_title = context_data.get("DOCUMENT_TITLE", "SAT_Report")
                    else:
                        doc_title = "SAT_Report"
                else:
                    doc_title = "SAT_Report"
            except Exception as db_error:
                current_app.logger.warning(f"Database error when getting title, using default: {db_error}")
                doc_title = "SAT_Report"
            
            # Get project number for filename (SAT_PROJNUMBER format)
            project_number = context_data.get("PROJECT_REFERENCE", "").strip()
            if not project_number:
                project_number = context_data.get("PROJECT_NUMBER", "").strip()
            if not project_number:
                project_number = submission_id[:8]  # Fallback to submission ID
                
            # Clean project number for filename  
            safe_proj_num = "".join(c if c.isalnum() or c in ['_', '-'] else "_" for c in project_number)
            download_name = f"SAT_{safe_proj_num}.docx"
            
            # Ensure file is not corrupted and has proper headers
            if not os.path.exists(permanent_path) or os.path.getsize(permanent_path) < 1000:
                current_app.logger.error(f"Existing file is corrupted or too small: {permanent_path}")
                flash('Report file is corrupted. Please regenerate.', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))
            
            current_app.logger.info(f"Serving existing file: {permanent_path} as {download_name}")
            
            # Return with proper Word document headers
            return send_file(
                permanent_path, 
                as_attachment=True, 
                download_name=download_name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        # If file doesn't exist, try to get data from database and generate
        try:
            from models import Report, SATReport
            report = Report.query.filter_by(id=submission_id).first()
            if not report:
                current_app.logger.error(f"Report not found in database: {submission_id}")
                flash('Report not found in database.', 'error')
                return redirect(url_for('dashboard.home'))

            sat_report = SATReport.query.filter_by(report_id=submission_id).first()
            if not sat_report:
                current_app.logger.error(f"SAT report data not found: {submission_id}")
                flash('Report data not found.', 'error')
                return redirect(url_for('dashboard.home'))

            # Parse stored data
            try:
                stored_data = json.loads(sat_report.data_json) if sat_report.data_json else {}
            except json.JSONDecodeError as json_error:
                current_app.logger.error(f"JSON decode error: {json_error}")
                stored_data = {}

            context_data = stored_data.get("context", {})
            if not context_data:
                current_app.logger.error(f"No context data found for submission: {submission_id}")
                flash('No report data found.', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))

        except Exception as db_error:
            current_app.logger.error(f"Database error: {db_error}")
            flash('Database connection error. Cannot generate report.', 'error')
            return redirect(url_for('dashboard.home'))

        # Generate fresh report
        current_app.logger.info(f"Generating fresh report for submission {submission_id}")

        try:
            # Check template file exists
            template_file = current_app.config.get('TEMPLATE_FILE', 'templates/SAT_Template.docx')
            if not os.path.exists(template_file):
                current_app.logger.error(f"Template file not found: {template_file}")
                flash('Report template file not found.', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))

            # PRESERVE EXACT TEMPLATE FORMAT - Open original and replace content only
            from docx import Document
            import re
            
            # Open the original SAT_Template.docx to preserve ALL formatting
            doc = Document(template_file)
            current_app.logger.info(f"Opened original SAT_Template.docx to preserve exact formatting: {template_file}")
            
            # AGGRESSIVE DEBUG: Print EVERYTHING
            current_app.logger.info("="*50)
            current_app.logger.info("FULL CONTEXT_DATA DUMP:")
            for key, value in context_data.items():
                current_app.logger.info(f"  '{key}': '{value}'")
            current_app.logger.info("="*50)
            
            # Create comprehensive mapping with more field variations
            replacement_data = {
                'DOCUMENT_TITLE': (
                    context_data.get('DOCUMENT_TITLE') or 
                    context_data.get('document_title') or 
                    context_data.get('Document_Title') or 
                    context_data.get('documentTitle') or ''
                ),
                'PROJECT_REFERENCE': (
                    context_data.get('PROJECT_REFERENCE') or 
                    context_data.get('project_reference') or 
                    context_data.get('Project_Reference') or ''
                ),
                'DOCUMENT_REFERENCE': (
                    context_data.get('DOCUMENT_REFERENCE') or 
                    context_data.get('document_reference') or 
                    context_data.get('Document_Reference') or 
                    context_data.get('doc_reference') or ''
                ),
                'DATE': (
                    context_data.get('DATE') or 
                    context_data.get('date') or 
                    context_data.get('Date') or ''
                ),
                'CLIENT_NAME': (
                    context_data.get('CLIENT_NAME') or 
                    context_data.get('client_name') or 
                    context_data.get('Client_Name') or ''
                ),
                'REVISION': (
                    context_data.get('REVISION') or 
                    context_data.get('revision') or 
                    context_data.get('Revision') or 
                    context_data.get('rev') or ''
                ),
                'PREPARED_BY': context_data.get('PREPARED_BY', context_data.get('prepared_by', '')),
                'PREPARER_DATE': context_data.get('PREPARER_DATE', context_data.get('preparer_date', '')),
                'REVIEWED_BY_TECH_LEAD': context_data.get('REVIEWED_BY_TECH_LEAD', context_data.get('reviewed_by_tech_lead', '')),
                'TECH_LEAD_DATE': context_data.get('TECH_LEAD_DATE', context_data.get('tech_lead_date', '')),
                'REVIEWED_BY_PM': context_data.get('REVIEWED_BY_PM', context_data.get('reviewed_by_pm', '')),
                'PM_DATE': context_data.get('PM_DATE', context_data.get('pm_date', '')),
                'APPROVED_BY_CLIENT': context_data.get('APPROVED_BY_CLIENT', context_data.get('approved_by_client', '')),
                'PURPOSE': context_data.get('PURPOSE', context_data.get('purpose', '')),
                'SCOPE': context_data.get('SCOPE', context_data.get('scope', '')),
                'REVISION_DETAILS': context_data.get('REVISION_DETAILS', context_data.get('revision_details', '')),
                'REVISION_DATE': context_data.get('REVISION_DATE', context_data.get('revision_date', '')),
                # Add signature placeholders
                'SIG_PREPARED': '',
                'SIG_REVIEW_TECH': '',
                'SIG_REVIEW_PM': '',
                'SIG_APPROVAL_CLIENT': ''
            }
            
            # Log final values for debugging
            current_app.logger.info(f"Final DOCUMENT_TITLE value: '{replacement_data['DOCUMENT_TITLE']}'")
            current_app.logger.info(f"Final DOCUMENT_REFERENCE value: '{replacement_data['DOCUMENT_REFERENCE']}'")
            current_app.logger.info(f"Final REVISION value: '{replacement_data['REVISION']}'")
            current_app.logger.info(f"Final PROJECT_REFERENCE value: '{replacement_data['PROJECT_REFERENCE']}'")
            
            def clean_text(text):
                """Clean template text by first removing Jinja2, then replacing tags"""
                if not text.strip():
                    return text
                
                original_text = text
                import re
                
                # FIRST: Remove all Jinja2 template syntax BEFORE doing replacements
                # Remove {% for %} ... {% endfor %} blocks (most common issue)
                text = re.sub(r'{%\s*for\s+[^%]*%}.*?{%\s*endfor\s*%}', '', text, flags=re.DOTALL)
                
                # Remove standalone {% %} blocks
                text = re.sub(r'{%\s*endfor\s*%}', '', text)
                text = re.sub(r'{%\s*for\s+[^%]*%}', '', text)
                text = re.sub(r'{%\s*[^%]*\s*%}', '', text)
                
                # SECOND: Replace template tags with actual values
                for tag, value in replacement_data.items():
                    if value:  # Only replace if we have a value
                        patterns_to_try = [
                            f'{{{{ {tag} }}}}',     # {{ TAG }}
                            f'{{{{{tag}}}}}',       # {{TAG}}
                            f'{{{{  {tag}  }}}}',   # {{  TAG  }}
                            f'{{{{ {tag}}}}}',      # {{ TAG}}
                            f'{{{{{tag} }}}}',      # {{TAG }}
                        ]
                        for pattern in patterns_to_try:
                            if pattern in text:
                                text = text.replace(pattern, str(value))
                                current_app.logger.info(f"REPLACED '{pattern}' with '{value}' in text")
                
                # THIRD: Remove any remaining empty template tags
                text = re.sub(r'{{\s*[^}]*\s*}}', '', text)
                
                # FOURTH: Clean up extra whitespace
                text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
                text = text.strip()
                
                # Log results
                if '{{' in text and text != original_text:
                    current_app.logger.info(f"STILL HAS TAGS after replacement: '{text[:100]}...'")
                elif '{{' in original_text and '{{' not in text:
                    current_app.logger.info(f"SUCCESSFULLY CLEANED: '{original_text[:50]}...' -> '{text[:50]}...'")
                
                return text
            
            # Advanced replacement: Process runs within paragraphs (handles split template tags)
            def replace_in_runs(paragraph):
                """Replace template tags that might be split across runs in Word"""
                if not paragraph.runs:
                    return False
                
                # Get full text from all runs
                full_text = ''.join(run.text for run in paragraph.runs)
                if not full_text or ('{{' not in full_text and '{%' not in full_text):
                    return False
                
                # Debug: Show what we found
                if any(tag in full_text for tag in ['DOCUMENT_TITLE', 'DOCUMENT_REFERENCE', 'REVISION']):
                    current_app.logger.info(f"FOUND TARGET TAG: '{full_text[:100]}...'")
                
                # Clean the full text
                new_full_text = clean_text(full_text)
                if new_full_text == full_text:
                    return False
                
                current_app.logger.info(f"RUN REPLACEMENT: '{full_text[:50]}...' -> '{new_full_text[:50]}...'")
                
                # Clear all runs and add new text as single run
                for run in paragraph.runs:
                    run.clear()
                
                # Add the cleaned text as a new run
                if new_full_text.strip():
                    paragraph.add_run(new_full_text)
                
                return True
            
            # Replace text in paragraphs (preserves all formatting)
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    # Try run-level replacement first
                    if replace_in_runs(paragraph):
                        current_app.logger.info(f"Para {paragraph_count}: RUNS PROCESSED")
                    elif '{{' in paragraph.text:
                        current_app.logger.info(f"Para {paragraph_count} UNCHANGED but has template tags: '{paragraph.text[:50]}...'")
                    paragraph_count += 1
            
            # Replace text in tables (preserves table formatting)
            table_count = 0
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if paragraph.text and paragraph.text.strip():
                                # Try run-level replacement for table cells too
                                if replace_in_runs(paragraph):
                                    current_app.logger.info(f"Table {table_count} Row {row_idx} Cell {cell_idx}: RUNS PROCESSED")
                                elif '{{' in paragraph.text:
                                    current_app.logger.info(f"Table {table_count} Row {row_idx} Cell {cell_idx} UNCHANGED: '{paragraph.text[:30]}...'")
                table_count += 1
            
            # Process headers and footers (Word documents can have these)
            header_footer_count = 0
            for section in doc.sections:
                # Process headers
                if hasattr(section, 'header'):
                    for paragraph in section.header.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            if replace_in_runs(paragraph):
                                current_app.logger.info(f"Header {header_footer_count}: RUNS PROCESSED")
                            elif '{{' in paragraph.text:
                                current_app.logger.info(f"Header {header_footer_count} UNCHANGED: '{paragraph.text[:30]}...'")
                
                # Process footers
                if hasattr(section, 'footer'):
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            if replace_in_runs(paragraph):
                                current_app.logger.info(f"Footer {header_footer_count}: RUNS PROCESSED")
                            elif '{{' in paragraph.text:
                                current_app.logger.info(f"Footer {header_footer_count} UNCHANGED: '{paragraph.text[:30]}...'")
                
                header_footer_count += 1
            
            current_app.logger.info("Replaced template tags in document, tables, headers, and footers")

            # Render template with field tags using FIXED approach
            try:
                # Ensure output directory exists
                permanent_dir = current_app.config['OUTPUT_DIR']
                os.makedirs(permanent_dir, exist_ok=True)
                
                # Template content already exists - no need to add sections
                # All formatting, logos, headers, footers, styles are preserved
                current_app.logger.info("Original template structure and formatting preserved")
                
                # Save using FIXED approach (memory buffer to avoid corruption)
                try:
                    import io
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer_size = len(buffer.getvalue())
                    current_app.logger.info(f"Template document saved to memory buffer: {buffer_size} bytes")
                    
                    # Write to file using working method
                    buffer.seek(0)
                    with open(permanent_path, 'wb') as f:
                        f.write(buffer.getvalue())
                    
                    current_app.logger.info(f"SAT template document written to file: {permanent_path}")
                    
                except Exception as save_error:
                    current_app.logger.error(f"Template save failed: {save_error}")
                    raise Exception(f"Failed to save template document: {save_error}")
                
                # Verify file was created and has reasonable size
                if not os.path.exists(permanent_path):
                    raise Exception("Document file was not created")
                    
                file_size = os.path.getsize(permanent_path)
                if file_size < 1000:  # Word docs should be at least 1KB
                    raise Exception(f"Document file too small ({file_size} bytes) - likely corrupted")
                    
                current_app.logger.info(f"Document verified: {permanent_path} ({file_size} bytes)")
                
                # Verify the file was created and has content
                if not os.path.exists(permanent_path) or os.path.getsize(permanent_path) == 0:
                    raise Exception("Document file was not created properly or is empty")
                    
                current_app.logger.info(f"Document saved successfully: {permanent_path} ({os.path.getsize(permanent_path)} bytes)")
                
            except Exception as render_error:
                current_app.logger.error(f"Error rendering/saving document: {render_error}", exc_info=True)
                flash(f'Error generating report document: {str(render_error)}', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))

            current_app.logger.info(f"Fresh report generated: {permanent_path}")

            # Get project number for filename (SAT_PROJNUMBER format)
            project_number = context_data.get("PROJECT_REFERENCE", "").strip()
            if not project_number:
                project_number = context_data.get("PROJECT_NUMBER", "").strip()
            if not project_number:
                project_number = submission_id[:8]  # Fallback to submission ID
                
            # Clean project number for filename
            safe_proj_num = "".join(c if c.isalnum() or c in ['_', '-'] else "_" for c in project_number)
            download_name = f"SAT_{safe_proj_num}.docx"

            # Verify file exists and has proper size before sending
            if not os.path.exists(permanent_path) or os.path.getsize(permanent_path) == 0:
                flash('Error: Generated document is empty or corrupted.', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))

            current_app.logger.info(f"Serving file: {permanent_path} as {download_name}")
            
            # Test different download approaches
            current_app.logger.info(f"Testing direct file serve without modifications")
            
            # First verify the file on server is good
            try:
                from docx import Document
                test_doc = Document(permanent_path)
                para_count = len(test_doc.paragraphs)
                current_app.logger.info(f"Server verification: Document has {para_count} paragraphs and can be opened")
            except Exception as verify_error:
                current_app.logger.error(f"Document corrupt on server: {verify_error}")
                flash('Document is corrupted on server', 'error')
                return redirect(url_for('status.view_status', submission_id=submission_id))
            
            # Try serving the file with minimal processing
            try:
                # Read file into memory and serve from memory to avoid file locking issues
                with open(permanent_path, 'rb') as f:
                    file_data = f.read()
                
                current_app.logger.info(f"Read {len(file_data)} bytes from file")
                
                # Create response from memory
                from flask import Response
                response = Response(
                    file_data,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={
                        'Content-Disposition': f'attachment; filename="{download_name}"',
                        'Content-Length': str(len(file_data))
                    }
                )
                
                current_app.logger.info(f"Serving {download_name} from memory ({len(file_data)} bytes)")
                return response
                
            except Exception as serve_error:
                current_app.logger.error(f"Error serving from memory: {serve_error}")
                # Final fallback
                return send_file(permanent_path, as_attachment=True, download_name=download_name)

        except Exception as generation_error:
            current_app.logger.error(f"Error during report generation: {generation_error}", exc_info=True)
            flash('Error generating report for download.', 'error')
            return redirect(url_for('status.view_status', submission_id=submission_id))

    except Exception as e:
        current_app.logger.error(f"Error in download_report for {submission_id}: {e}", exc_info=True)
        flash('Error downloading report.', 'error')
        return redirect(url_for('dashboard.home'))



@status_bp.route('/list')
@login_required
def list_submissions():
    """List all submissions for admin view"""
    from models import Report, SATReport

    try:
        reports = Report.query.order_by(Report.created_at.desc()).all()
        submission_list = []

        for report in reports:
            sat_report = SATReport.query.filter_by(report_id=report.id).first()
            if not sat_report:
                continue

            try:
                stored_data = json.loads(sat_report.data_json)
            except json.JSONDecodeError:
                stored_data = {}

            # Determine overall status
            if report.approvals_json:
                try:
                    approvals = json.loads(report.approvals_json)
                    statuses = [a.get("status", "pending") for a in approvals]
                    if "rejected" in statuses:
                        overall_status = "rejected"
                    elif all(status == "approved" for status in statuses):
                        overall_status = "approved"
                    elif any(status == "approved" for status in statuses):
                        overall_status = "partially_approved"
                    else:
                        overall_status = "pending"
                except:
                    overall_status = "pending"
            else:
                overall_status = "draft"

            submission_list.append({
                "id": report.id,
                "document_title": stored_data.get("context", {}).get("DOCUMENT_TITLE", "SAT Report"),
                "client_name": stored_data.get("context", {}).get("CLIENT_NAME", ""),
                "created_at": report.created_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.created_at, dt.datetime) else report.created_at,
                "updated_at": report.updated_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(report.updated_at, dt.datetime) else report.updated_at,
                "status": overall_status,
                "user_email": report.user_email
            })

        return render_template('submissions_list.html', submissions=submission_list)

    except Exception as e:
        current_app.logger.error(f"Error fetching submissions list: {e}")
        flash('Error loading submissions.', 'error')
        return render_template('submissions_list.html', submissions=[])